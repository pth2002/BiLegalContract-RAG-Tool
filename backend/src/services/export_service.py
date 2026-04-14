"""Export service for generating reports."""

from datetime import datetime
from io import BytesIO
from typing import Optional

from docx import Document as DocxDocument 
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ..models import AnalysisResult, ExportFormat, RiskCard, Document, DocumentAnalysis


class ExportError(Exception):
    """Exception raised during export."""

    pass


def generate_markdown_report(
    document: Document,
    analysis: DocumentAnalysis,
    include_risks: bool = True,
    include_summary: bool = True,
) -> str:
    """Generate Markdown format report.

    Args:
        document: Document being analyzed
        analysis: Analysis results
        include_risks: Include risk details
        include_summary: Include executive summary

    Returns:
        Markdown formatted report
    """
    perspective_name = "甲方视角" if analysis.perspective.value == "party_a" else "乙方视角"
    risks = analysis.risks

    lines = [
        "# 合同智能审查报告",
        "",
        f"**生成时间**: {analysis.analyzed_at.strftime('%Y-%m-%d %H:%M:%S')}",
        f"**审查视角**: {perspective_name}",
        f"**文档**: {document.filename}",
        f"**耗时**: {analysis.duration_ms / 1000:.1f} 秒",
        "",
        "---",
        "",
    ]

    if include_summary:
        # Count risks by severity
        high = sum(1 for r in risks if r.severity.value == "高")
        medium = sum(1 for r in risks if r.severity.value == "中")
        low = sum(1 for r in risks if r.severity.value == "低")

        lines.extend([
            "## 概览",
            "",
            f"本次审查共发现 **{len(risks)} 个风险点**，其中：",
            f"- 高风险：{high} 个",
            f"- 中风险：{medium} 个",
            f"- 低风险：{low} 个",
            "",
            "---",
            "",
        ])

    if include_risks and risks:
        lines.extend([
            "## 风险详情",
            "",
        ])

        # Group risks by severity
        severity_order = ["高", "中", "低"]
        for severity in severity_order:
            severity_risks = [r for r in risks if r.severity.value == severity]

            if severity_risks:
                lines.append(f"### {severity}风险")
                lines.append("")

                for i, risk in enumerate(severity_risks, 1):
                    lines.extend([
                        f"#### {i}. {risk.clause_title}（{risk.severity.value}）",
                        "",
                        "**原文**：",
                        f"> {risk.original_text}",
                        "",
                        "**风险描述**：",
                        risk.risk_description,
                        "",
                        "**修改建议**：",
                        risk.suggested_revision,
                        "",
                        "---",
                        "",
                    ])

    if include_summary:
        lines.extend([
            "## 总结",
            "",
            analysis.summary,
            "",
        ])

    return "\n".join(lines)


def generate_markdown_download(
    document: Document,
    analysis: DocumentAnalysis,
    include_risks: bool = True,
    include_summary: bool = True,
) -> tuple[bytes, str]:
    """Generate Markdown report for download.

    Args:
        document: Document being analyzed
        analysis: Analysis results
        include_risks: Include risk details
        include_summary: Include executive summary

    Returns:
        Tuple of (file content bytes, suggested filename)
    """
    content = generate_markdown_report(
        document=document,
        analysis=analysis,
        include_risks=include_risks,
        include_summary=include_summary,
    )

    timestamp = datetime.now().strftime("%Y%m%d")
    download_name = f"合同审查报告_{timestamp}.md"

    return content.encode("utf-8"), download_name


def generate_docx_report(
    document: Document,
    analysis: DocumentAnalysis,
    include_risks: bool = True,
    include_summary: bool = True,
) -> BytesIO:
    """Generate DOCX format report.

    Args:
        document: Document being analyzed
        analysis: Analysis results
        include_risks: Include risk details
        include_summary: Include executive summary

    Returns:
        BytesIO containing the DOCX file
    """
    perspective_name = "甲方视角" if analysis.perspective.value == "party_a" else "乙方视角"
    risks = analysis.risks

    doc = DocxDocument()

    # Title
    title = doc.add_heading('合同智能审查报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Metadata
    doc.add_paragraph(f"生成时间: {analysis.analyzed_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"审查视角: {perspective_name}")
    doc.add_paragraph(f"文档: {document.filename}")
    doc.add_paragraph(f"耗时: {analysis.duration_ms / 1000:.1f} 秒")
    doc.add_paragraph()

    if include_summary:
        # Overview
        doc.add_heading('概览', level=1)

        # Count risks by severity
        high = sum(1 for r in risks if r.severity.value == "高")
        medium = sum(1 for r in risks if r.severity.value == "中")
        low = sum(1 for r in risks if r.severity.value == "低")

        doc.add_paragraph(f"本次审查共发现 {len(risks)} 个风险点，其中：")
        doc.add_paragraph(f"高风险：{high} 个", style='List Bullet')
        doc.add_paragraph(f"中风险：{medium} 个", style='List Bullet')
        doc.add_paragraph(f"低风险：{low} 个", style='List Bullet')
        doc.add_paragraph()

    if include_risks and risks:
        # Risk details
        doc.add_heading('风险详情', level=1)

        # Group risks by severity
        severity_order = [("高", "高风险"), ("中", "中风险"), ("低", "低风险")]

        for severity_value, severity_label in severity_order:
            severity_risks = [r for r in risks if r.severity.value == severity_value]

            if severity_risks:
                doc.add_heading(severity_label, level=2)

                for i, risk in enumerate(severity_risks, 1):
                    # Risk title with clause and severity
                    risk_title = doc.add_heading(f"{i}. {risk.clause_title}", level=3)

                    # Original text
                    p = doc.add_paragraph()
                    p.add_run("原文：").bold = True
                    doc.add_paragraph(risk.original_text, style='Quote')

                    # Risk description
                    p = doc.add_paragraph()
                    p.add_run("风险描述：").bold = True
                    doc.add_paragraph(risk.risk_description)

                    # Suggested revision
                    p = doc.add_paragraph()
                    p.add_run("修改建议：").bold = True
                    suggestion_para = doc.add_paragraph(risk.suggested_revision)
                    suggestion_para.paragraph_format.left_indent = Inches(0.25)

                    doc.add_paragraph()  # spacing between risks

    if include_summary:
        # Summary
        doc.add_heading('总结', level=1)
        doc.add_paragraph(analysis.summary)

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output
